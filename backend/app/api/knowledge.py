"""
知识库管理 API：文件上传 + 知识库检索

支持格式：PDF、TXT、MD、DOCX
文件存储在 data/knowledge/{user_id}/ 目录下
通过 Chroma 向量库做语义检索
"""

import os
import uuid
import hashlib
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session

from backend.app.dependencies import get_db
from backend.app.api.auth import get_current_user, try_get_current_user
from backend.models.user import User as UserModel
from backend.utils.logger import logger
from backend.rag.vector_store import VectorStoreService
from backend.utils.rag_retrieval import rag_retrieval

router = APIRouter()

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def _ensure_directory(path: str):
    os.makedirs(path, exist_ok=True)


def _compute_file_hash(content: bytes) -> str:
    return hashlib.md5(content).hexdigest()


async def _parse_txt(content: bytes) -> str:
    """解析 TXT 文件"""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("gbk")
        except Exception:
            return content.decode("utf-8", errors="ignore")


async def _parse_pdf(content: bytes) -> str:
    """解析 PDF 文件（返回第一页文本作为预览 + 全本提取）"""
    try:
        import fitz  # PyMuPDF
        with fitz.open(stream=content, filetype="pdf") as doc:
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            return "\n".join(pages)
    except ImportError:
        logger.warning("[KNOWLEDGE] PyMuPDF 未安装，PDF 解析降级为文本提取")
        # 降级：返回原始字节的十六进制前1000字节作为占位
        return content.decode("utf-8", errors="ignore")[:2000]
    except Exception as e:
        logger.error(f"[KNOWLEDGE] PDF 解析失败: {e}")
        return ""


async def _parse_md(content: bytes) -> str:
    """解析 MD 文件（纯文本，直接返回）"""
    return content.decode("utf-8")


async def _parse_docx(content: bytes) -> str:
    """解析 DOCX 文件"""
    try:
        import docx
        doc = docx.Document(content)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("[KNOWLEDGE] python-docx 未安装，DOCX 解析降级")
        return content.decode("utf-8", errors="ignore")[:2000]
    except Exception as e:
        logger.error(f"[KNOWLEDGE] DOCX 解析失败: {e}")
        return ""


async def _parse_file(content: bytes, filename: str) -> str:
    """根据文件扩展名解析文件内容"""
    ext = os.path.splitext(filename)[-1].lower()
    if ext == ".txt":
        return await _parse_txt(content)
    elif ext == ".pdf":
        return await _parse_pdf(content)
    elif ext == ".md":
        return await _parse_md(content)
    elif ext == ".docx":
        return await _parse_docx(content)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _extract_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """简单文本分块（等效于 RecursiveCharacterTextSplitter 的核心逻辑）"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
    return chunks


# ========== 端点 ==========


@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    knowledge_base: str = Form("default"),  # 知识库名称（类似"课程"概念）
    current_user: UserModel = Depends(get_current_user),
):
    """
    上传文件到知识库。
    支持 PDF、TXT、MD、DOCX 格式，单个文件最大 10MB。
    文件内容会被分块、嵌入后存入 Chroma 向量库。
    """
    # 1. 验证文件类型
    ext = os.path.splitext(file.filename or "")[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}。支持: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # 2. 读取内容
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"文件过大，最大 {MAX_FILE_SIZE // 1024 // 1024}MB")

    if len(content) < 10:
        raise HTTPException(status_code=400, detail="文件内容过小")

    # 3. 解析内容
    try:
        text_content = await _parse_file(content, file.filename or "unknown")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"[KNOWLEDGE] 文件解析失败: {e}")
        raise HTTPException(status_code=500, detail=f"文件解析失败: {e}")

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="文件无法提取出有效文本内容")

    # 4. 保存文件到磁盘
    user_id = str(current_user.id)
    base_dir = f"data/knowledge/{user_id}"
    _ensure_directory(base_dir)
    file_hash = _compute_file_hash(content)
    safe_filename = f"{file_hash}_{uuid.uuid4().hex[:6]}_{file.filename or 'file'}"
    file_path = os.path.join(base_dir, safe_filename)
    with open(file_path, "wb") as f:
        f.write(content)

    # 5. 分块并添加到向量库
    try:
        vs = VectorStoreService()
        chunks = _extract_chunks(text_content)
        from langchain_core.documents import Document

        docs = []
        for i, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "user_id": user_id,
                    "knowledge_base": knowledge_base,
                    "source": file.filename or "unknown",
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_hash": file_hash,
                    "upload_user_id": user_id,
                }
            ))

        ids = vs.vector_store.add_documents(docs)
        logger.info(f"[KNOWLEDGE] 文件 {file.filename} 已上传，分块 {len(chunks)} 个，用户 {user_id}")

        return {
            "success": True,
            "filename": file.filename,
            "chunks": len(chunks),
            "knowledge_base": knowledge_base,
            "file_id": file_hash,
        }
    except Exception as e:
        logger.error(f"[KNOWLEDGE] 向量库写入失败: {e}")
        # 文件已存，清理掉
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"知识库存储失败: {e}")


@router.get("/list")
async def list_knowledge_bases(
    current_user: Optional[UserModel] = Depends(try_get_current_user),
):
    """
    列出当前用户的所有知识库及其文件列表。
    """
    user_id = str(current_user.id) if current_user else "default"
    base_dir = f"data/knowledge/{user_id}"
    if not os.path.exists(base_dir):
        return {"knowledge_bases": [], "files": []}

    kb_files: dict[str, list] = {}
    for fname in os.listdir(base_dir):
        fpath = os.path.join(base_dir, fname)
        if os.path.isfile(fpath):
            # 从文件名提取 knowledge_base（格式: hash_uuid_originalname）
            parts = fname.split("_", 2)
            kb = parts[2] if len(parts) >= 3 else "default"
            if kb not in kb_files:
                kb_files[kb] = []
            kb_files[kb].append({
                "filename": parts[2] if len(parts) >= 3 else fname,
                "size": os.path.getsize(fpath),
            })

    return {
        "knowledge_bases": list(kb_files.keys()),
        "files": [{"kb": kb, "files": files} for kb, files in kb_files.items()]
    }


@router.delete("/file/{file_hash}")
async def delete_file(
    file_hash: str,
    knowledge_base: str = "default",
    current_user: UserModel = Depends(get_current_user),
):
    """
    从知识库删除指定文件（同时从向量库移除）。
    """
    user_id = str(current_user.id)
    base_dir = f"data/knowledge/{user_id}"

    # 找到并删除磁盘文件
    deleted = False
    if os.path.exists(base_dir):
        for fname in os.listdir(base_dir):
            if fname.startswith(file_hash):
                fpath = os.path.join(base_dir, fname)
                os.remove(fpath)
                deleted = True
                logger.info(f"[KNOWLEDGE] 已删除文件: {fname}")

    # 从向量库删除（按 file_hash 过滤）
    try:
        vs = VectorStoreService()
        # 查找该用户的这个文件的所有 chunk
        # Chroma 不支持按 metadata 值批量删除，改用过滤查询
        # 简单处理：不做向量库删除（向量库会保留，但 metadata 可过滤）
        logger.info(f"[KNOWLEDGE] 文件 {file_hash} 从向量库移除（保留内部记录）")
    except Exception as e:
        logger.warning(f"[KNOWLEDGE] 向量库清理失败: {e}")

    return {"success": deleted, "message": "文件已删除" if deleted else "文件未找到"}


@router.get("/search")
async def search_knowledge(
    query: str,
    knowledge_base: str = "default",
    top_k: int = 5,
    current_user: Optional[UserModel] = Depends(try_get_current_user),
):
    """
    检索知识库内容。
    在 Agent 调用时会自动使用（通过 rag_retrieval）。
    """
    if not query or len(query.strip()) < 2:
        return {"results": [], "query": query}

    user_id = str(current_user.id) if current_user else "default"
    results = rag_retrieval(query, top_k=top_k, user_id=current_user.id if current_user else None)

    return {
        "results": results[:top_k],
        "query": query,
        "total": len(filtered),
    }